# core/docgen.py
from __future__ import annotations
from pathlib import Path
from typing import Dict
from docx import Document

def _replace_in_paragraphs(doc, mapping: Dict[str, str]) -> None:
    for p in doc.paragraphs:
        for key, val in mapping.items():
            if key in p.text:
                # reconstruir runs para no perder formato básico
                inline = p.runs
                text = "".join(run.text for run in inline)
                text = text.replace(key, val)
                # vaciar y rearmar en un solo run
                for idx in range(len(inline)-1, -1, -1):
                    p.runs[idx].clear()
                    try:
                        p._element.remove(p.runs[idx]._element)
                    except Exception:
                        pass
                p.add_run(text)

def _replace_in_tables(doc, mapping: Dict[str, str]) -> None:
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in mapping.items():
                    if key in cell.text:
                        # Reemplazos en cada párrafo de la celda
                        for p in cell.paragraphs:
                            if key in p.text:
                                inline = p.runs
                                text = "".join(run.text for run in inline)
                                text = text.replace(key, val)
                                for idx in range(len(inline)-1, -1, -1):
                                    p.runs[idx].clear()
                                    try:
                                        p._element.remove(p.runs[idx]._element)
                                    except Exception:
                                        pass
                                p.add_run(text)

def render_docx(template_path: Path, output_path: Path, mapping: Dict[str, str]) -> Path:
    """
    Carga un DOCX de plantilla y sustituye los placeholders indicados en 'mapping'.
    Devuelve la ruta resultante.
    *Nota:* Los placeholders deben estar en una sola 'run' en Word para que el reemplazo sea confiable.
    """
    doc = Document(str(template_path))
    _replace_in_paragraphs(doc, mapping)
    _replace_in_tables(doc, mapping)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path